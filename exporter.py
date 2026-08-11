import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from models import Question, QuestionType


def add_header_table(doc: docx.Document, header_info: dict, test_code: str = ""):
    """Chèn bảng tiêu đề 2x2 định dạng chuẩn đề thi GDPT vào đầu trang."""
    if not header_info:
        return

    school_name = str(header_info.get("school_name", "")).strip().upper()
    exam_title = str(header_info.get("exam_title", "")).strip().upper()
    sub_title = str(header_info.get("sub_title", "")).strip()
    duration = header_info.get("duration", 90)

    if test_code and f"Mã đề: {test_code}" not in sub_title:
        sub_title = f"{sub_title} - Mã đề: {test_code}" if sub_title else f"Mã đề: {test_code}"

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(3.8)

    # Ô (0,0): Tên trường / Trung tâm
    p00 = table.cell(0, 0).paragraphs[0]
    p00.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r00 = p00.add_run(school_name)
    r00.bold = True
    r00.font.name = "Times New Roman"
    r00.font.size = Pt(11)

    # Ô (0,1): Tên đề / Kỳ thi
    p01 = table.cell(0, 1).paragraphs[0]
    p01.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r01 = p01.add_run(exam_title)
    r01.bold = True
    r01.font.name = "Times New Roman"
    r01.font.size = Pt(11)

    # Ô (1,0): Ghi chú / Môn học / Mã đề
    p10 = table.cell(1, 0).paragraphs[0]
    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r10 = p10.add_run(sub_title)
    r10.font.name = "Times New Roman"
    r10.font.size = Pt(10.5)

    # Ô (1,1): Thời gian làm bài
    p11 = table.cell(1, 1).paragraphs[0]
    p11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r11 = p11.add_run(f"Thời gian: {duration} phút (không kể thời gian phát đề)")
    r11.font.name = "Times New Roman"
    r11.font.size = Pt(10.5)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(12)
    r_line = p_line.add_run("────────────────────────────────────────")
    r_line.font.name = "Times New Roman"
    r_line.font.size = Pt(9)


def safe_add_image(doc: docx.Document, img_path: str, max_width_inches: float = 5.0):
    """Chèn hình ảnh an toàn vào tài liệu nếu tệp tồn tại."""
    if not img_path or not isinstance(img_path, str):
        return
    if os.path.exists(img_path):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(4)
            p_img.paragraph_format.space_after = Pt(6)
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(max_width_inches))
        except Exception:
            pass


def create_combined_answer_docx(generated_exams_dict: dict, output_filepath: str):
    """Tạo file Bảng đáp án tổng hợp (Cột 1: Câu, các cột tiếp theo: Đáp án theo từng Mã đề)."""
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("BẢNG ĐÁP ÁN TỔNG HỢP CÁC MÃ ĐỀ THI")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()

    codes = list(generated_exams_dict.keys())
    if not codes:
        doc.save(output_filepath)
        return

    max_q_count = max(len(qs) for qs in generated_exams_dict.values())
    table = doc.add_table(rows=max_q_count + 1, cols=len(codes) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Câu"
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[0].font.name = "Times New Roman"

    for col_idx, code in enumerate(codes, start=1):
        hdr_cells[col_idx].text = f"Mã đề {code}"
        hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[col_idx].paragraphs[0].runs[0].font.name = "Times New Roman"

    for q_idx in range(max_q_count):
        row_cells = table.rows[q_idx + 1].cells
        row_cells[0].text = str(q_idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.name = "Times New Roman"

        for col_idx, code in enumerate(codes, start=1):
            qs = generated_exams_dict[code]
            ans = qs[q_idx].answer if q_idx < len(qs) else ""
            row_cells[col_idx].text = str(ans or "")
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_cells[col_idx].paragraphs[0].runs:
                row_cells[col_idx].paragraphs[0].runs[0].font.name = "Times New Roman"

    doc.save(output_filepath)


def combine_docx_files(file_list: list, output_filepath: str):
    """Gộp nhiều file Word lời giải chi tiết thành 1 file duy nhất."""
    if not file_list:
        return
    master = docx.Document(file_list[0])
    for fpath in file_list[1:]:
        if os.path.exists(fpath):
            master.add_page_break()
            sub_doc = docx.Document(fpath)
            for element in sub_doc.element.body:
                master.element.body.append(element)
    master.save(output_filepath)


def export_questions_to_word(
    questions: list,
    output_filepath: str,
    mode: str = "de_goc",
    ds_table_format: bool = True,
    tln_box_format: bool = True,
    test_code: str = "",
    header_info: dict = None
):
    """Xuất danh sách câu hỏi ra file Word (.docx) theo các tùy chọn định dạng."""
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    if mode == "dap_an":
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(f"BẢNG ĐÁP ÁN NHANH - MÃ ĐỀ: {test_code}" if test_code else "BẢNG ĐÁP ÁN NHANH")
        r_title.bold = True
        r_title.font.name = "Times New Roman"
        r_title.font.size = Pt(14)
        doc.add_paragraph()

        tbl_ans = doc.add_table(rows=1, cols=2)
        tbl_ans.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_ans.style = 'Table Grid'
        
        hdr_cells = tbl_ans.rows[0].cells
        hdr_cells[0].text = "Câu"
        hdr_cells[1].text = "Đáp án"
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, q in enumerate(questions, start=1):
            row_cells = tbl_ans.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(q.answer or "")
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.save(output_filepath)
        return

    if header_info and mode in ["de_goc", "de_dong_chua"]:
        add_header_table(doc, header_info, test_code)

    if mode == "loi_giai_chi_tiet":
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(f"LỜI GIẢI CHI TIẾT - MÃ ĐỀ: {test_code}" if test_code else "LỜI GIẢI CHI TIẾT ĐỀ THI")
        r_title.bold = True
        r_title.font.name = "Times New Roman"
        r_title.font.size = Pt(14)
        doc.add_paragraph()

    for idx, q in enumerate(questions, start=1):
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(4)
        
        r_lbl = p_q.add_run(f"Câu {idx}. ")
        r_lbl.bold = True
        r_lbl.font.name = "Times New Roman"
        
        r_cnt = p_q.add_run(q.content)
        r_cnt.font.name = "Times New Roman"

        if q.image_path:
            safe_add_image(doc, q.image_path)

        if q.format == QuestionType.TN and q.options:
            p_opts = doc.add_paragraph()
            p_opts.paragraph_format.space_after = Pt(4)
            for key in ['A', 'B', 'C', 'D']:
                val = q.options.get(key, '')
                if val:
                    r_k = p_opts.add_run(f"  {key}. ")
                    r_k.bold = True
                    r_k.font.name = "Times New Roman"
                    r_v = p_opts.add_run(f"{val}    ")
                    r_v.font.name = "Times New Roman"

        elif q.format == QuestionType.DS and q.tf_statements:
            if ds_table_format and mode in ["de_goc", "de_dong_chua"]:
                tbl_ds = doc.add_table(rows=len(q.tf_statements) + 1, cols=3)
                tbl_ds.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl_ds.style = 'Table Grid'
                
                tbl_ds.rows[0].cells[0].text = "Mệnh đề / Ý hỏi"
                tbl_ds.rows[0].cells[1].text = "Đúng"
                tbl_ds.rows[0].cells[2].text = "Sai"
                
                for c in tbl_ds.rows[0].cells:
                    c.paragraphs[0].runs[0].font.bold = True
                    c.paragraphs[0].runs[0].font.name = "Times New Roman"
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                tbl_ds.columns[0].width = Inches(5.2)
                tbl_ds.columns[1].width = Inches(0.8)
                tbl_ds.columns[2].width = Inches(0.8)

                for st_idx, (stmt_text, _) in enumerate(q.tf_statements, start=1):
                    row_cells = tbl_ds.rows[st_idx].cells
                    row_cells[0].text = stmt_text
                    row_cells[1].text = "[   ]"
                    row_cells[2].text = "[   ]"
                    row_cells[0].paragraphs[0].runs[0].font.name = "Times New Roman"
                    row_cells[1].paragraphs[0].runs[0].font.name = "Times New Roman"
                    row_cells[2].paragraphs[0].runs[0].font.name = "Times New Roman"
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            else:
                for stmt_text, status in q.tf_statements:
                    p_st = doc.add_paragraph()
                    p_st.paragraph_format.space_after = Pt(2)
                    r_st_txt = p_st.add_run(f"  • {stmt_text} ")
                    r_st_txt.font.name = "Times New Roman"
                    if mode == "loi_giai_chi_tiet":
                        r_st = p_st.add_run(f"[{status}]")
                        r_st.bold = True
                        r_st.font.name = "Times New Roman"
                        r_st.font.color.rgb = RGBColor(184, 84, 63)

        elif q.format == QuestionType.TLN:
            if tln_box_format and mode in ["de_goc", "de_dong_chua"]:
                p_tln = doc.add_paragraph()
                p_tln.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_tln.paragraph_format.space_before = Pt(4)
                r_box_lbl = p_tln.add_run("Đáp số:  ")
                r_box_lbl.bold = True
                r_box_lbl.font.name = "Times New Roman"

                tbl_box = doc.add_table(rows=1, cols=5)
                tbl_box.alignment = WD_TABLE_ALIGNMENT.RIGHT
                tbl_box.style = 'Table Grid'
                for cell in tbl_box.rows[0].cells:
                    cell.width = Inches(0.32)
                    p_cell = cell.paragraphs[0]
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cell.add_run(" ")
                doc.add_paragraph()

        if mode == "loi_giai_chi_tiet":
            if q.solution:
                p_sol = doc.add_paragraph()
                p_sol.paragraph_format.space_before = Pt(4)
                r_sol_head = p_sol.add_run("Lời giải: ")
                r_sol_head.bold = True
                r_sol_head.font.name = "Times New Roman"
                r_sol_txt = p_sol.add_run(q.solution)
                r_sol_txt.font.name = "Times New Roman"

            sol_img = getattr(q, 'solution_image_path', None)
            if sol_img:
                safe_add_image(doc, sol_img)

        if mode == "de_dong_chua":
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(4)
            p_space.paragraph_format.space_after = Pt(4)
            for _ in range(3):
                p_dot = doc.add_paragraph()
                p_dot.paragraph_format.space_after = Pt(2)
                r_dot = p_dot.add_run(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ─ strong. minor support font.name = "Times New Roman strict"

    doc.save(output_filepath)
